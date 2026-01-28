"""
Word to Markdown Converter

This module converts Word documents (.docx) to Markdown format
with support for different documentation categories:
- Architecture Document
- Functional Document
- Technical Document
"""

import os
import re
import argparse
from pathlib import Path
from enum import Enum
from typing import Optional
from datetime import datetime

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    raise ImportError("Please install python-docx: pip install python-docx")


class DocumentCategory(Enum):
    """Document category types."""
    ARCHITECTURE = "architecture"
    FUNCTIONAL = "functional"
    TECHNICAL = "technical"


class WordToMarkdownConverter:
    """
    Converts Word documents to Markdown format.

    Supports three documentation categories:
    - Architecture Document: System design, components, diagrams
    - Functional Document: Features, requirements, user stories
    - Technical Document: Implementation details, APIs, configurations
    """

    def __init__(self, category: DocumentCategory):
        """
        Initialize the converter with a document category.

        Args:
            category: The documentation category type
        """
        self.category = category
        self.output_dir = self._get_output_directory()

    def _get_output_directory(self) -> Path:
        """Get the output directory based on document category."""
        base_dir = Path(__file__).parent.parent / "docs"
        return base_dir / self.category.value

    def _convert_paragraph(self, paragraph: Paragraph) -> str:
        """
        Convert a Word paragraph to Markdown.

        Args:
            paragraph: The Word paragraph object

        Returns:
            Markdown formatted string
        """
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = paragraph.style.name.lower() if paragraph.style else ""

        # Handle headings
        if "heading 1" in style_name:
            return f"# {text}\n"
        elif "heading 2" in style_name:
            return f"## {text}\n"
        elif "heading 3" in style_name:
            return f"### {text}\n"
        elif "heading 4" in style_name:
            return f"#### {text}\n"
        elif "heading 5" in style_name:
            return f"##### {text}\n"
        elif "heading 6" in style_name:
            return f"###### {text}\n"
        elif "title" in style_name:
            return f"# {text}\n"
        elif "subtitle" in style_name:
            return f"## {text}\n"

        # Handle list items
        if "list" in style_name or "bullet" in style_name:
            return f"- {text}\n"
        if "number" in style_name:
            return f"1. {text}\n"

        # Handle inline formatting
        markdown_text = self._apply_inline_formatting(paragraph)

        return f"{markdown_text}\n"

    def _apply_inline_formatting(self, paragraph: Paragraph) -> str:
        """
        Apply inline formatting (bold, italic, code) to paragraph text.

        Args:
            paragraph: The Word paragraph object

        Returns:
            Markdown formatted string with inline styles
        """
        result = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            # Apply formatting
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            # Handle code style (monospace fonts)
            if run.font.name and "courier" in run.font.name.lower():
                text = f"`{text}`"

            result.append(text)

        return "".join(result) if result else paragraph.text

    def _convert_table(self, table: Table) -> str:
        """
        Convert a Word table to Markdown table format.

        Args:
            table: The Word table object

        Returns:
            Markdown formatted table string
        """
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

            # Add header separator after first row
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                rows.append(separator)

        return "\n".join(rows) + "\n"

    def _generate_metadata_header(self, filename: str) -> str:
        """
        Generate YAML front matter metadata for the Markdown file.

        Args:
            filename: Original Word document filename

        Returns:
            YAML front matter string
        """
        return f"""---
title: "{Path(filename).stem}"
category: "{self.category.value}"
type: "{self.category.name.replace('_', ' ').title()} Document"
created_date: "{datetime.now().strftime('%Y-%m-%d')}"
source_file: "{filename}"
---

"""

    def _get_category_template(self) -> str:
        """
        Get the template structure based on document category.

        Returns:
            Template string for the document category
        """
        templates = {
            DocumentCategory.ARCHITECTURE: """
> **Document Type**: Architecture Document
>
> This document describes the system architecture, components, and design decisions.

---

""",
            DocumentCategory.FUNCTIONAL: """
> **Document Type**: Functional Document
>
> This document describes the functional requirements, features, and user stories.

---

""",
            DocumentCategory.TECHNICAL: """
> **Document Type**: Technical Document
>
> This document describes the technical implementation, APIs, and configurations.

---

"""
        }
        return templates.get(self.category, "")

    def convert(self, input_path: str, output_filename: Optional[str] = None) -> str:
        """
        Convert a Word document to Markdown.

        Args:
            input_path: Path to the Word document (.docx)
            output_filename: Optional custom output filename

        Returns:
            Path to the generated Markdown file
        """
        input_path = Path(input_path)

        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")

        if input_path.suffix.lower() != ".docx":
            raise ValueError("Input file must be a .docx file")

        # Load the Word document
        doc = Document(str(input_path))

        # Build Markdown content
        markdown_parts = []

        # Add metadata header
        markdown_parts.append(self._generate_metadata_header(input_path.name))

        # Add category template
        markdown_parts.append(self._get_category_template())

        # Process document elements
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith("p"):
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        converted = self._convert_paragraph(paragraph)
                        if converted:
                            markdown_parts.append(converted)
                        break

            # Handle tables
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._tbl == element:
                        markdown_parts.append("\n")
                        markdown_parts.append(self._convert_table(table))
                        markdown_parts.append("\n")
                        break

        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Determine output filename
        if output_filename:
            output_name = output_filename if output_filename.endswith(".md") else f"{output_filename}.md"
        else:
            output_name = f"{input_path.stem}.md"

        # Clean filename
        output_name = re.sub(r'[^\w\-_.]', '_', output_name)

        output_path = self.output_dir / output_name

        # Write Markdown file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(markdown_parts))

        return str(output_path)


def main():
    """Main entry point for the converter CLI."""
    parser = argparse.ArgumentParser(
        description="Convert Word documents to Markdown",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Document Categories:
  architecture  - System design, components, diagrams
  functional    - Features, requirements, user stories
  technical     - Implementation details, APIs, configurations

Examples:
  python word_to_markdown.py input.docx -c architecture
  python word_to_markdown.py input.docx -c functional -o my_document
        """
    )

    parser.add_argument(
        "input",
        help="Path to the Word document (.docx)"
    )

    parser.add_argument(
        "-c", "--category",
        type=str,
        choices=["architecture", "functional", "technical"],
        required=True,
        help="Document category type"
    )

    parser.add_argument(
        "-o", "--output",
        type=str,
        default=None,
        help="Output filename (optional, defaults to input filename)"
    )

    args = parser.parse_args()

    # Map string to enum
    category_map = {
        "architecture": DocumentCategory.ARCHITECTURE,
        "functional": DocumentCategory.FUNCTIONAL,
        "technical": DocumentCategory.TECHNICAL
    }

    category = category_map[args.category]

    # Create converter and process document
    converter = WordToMarkdownConverter(category)

    try:
        output_path = converter.convert(args.input, args.output)
        print(f"Successfully converted document!")
        print(f"Output saved to: {output_path}")
    except FileNotFoundError as e:
        print(f"Error: {e}")
        exit(1)
    except ValueError as e:
        print(f"Error: {e}")
        exit(1)
    except Exception as e:
        print(f"Conversion failed: {e}")
        exit(1)


if __name__ == "__main__":
    main()
